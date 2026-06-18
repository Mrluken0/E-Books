import os
import re
import sys
import json
import glob
import argparse
import subprocess

# --- GESTION DES DÉPENDANCES ---
try:
    import docx
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
except ImportError:
    try:
        subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx", "--break-system-packages", "-q"])
        import docx
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
    except Exception as e:
        sys.stderr.write(f"Erreur d'installation de python-docx: {str(e)}\n")
        print(json.dumps({"status": "error", "message": f"Dependency error: {str(e)}"}))
        sys.exit(1)

# --- CONFIGURATION ENCODAGE ---
sys.stdout.reconfigure(encoding='utf-8')
sys.stderr.reconfigure(encoding='utf-8')

def ensure_toc_styles(doc):
    """Crée les styles 'TOC 1' et 'TOC 2' s'ils n'existent pas (basedOn Normal, indent gauche progressif)."""
    from docx.enum.style import WD_STYLE_TYPE
    styles = doc.styles
    base = styles['Normal']
    # (nom de style, indentation gauche en points)
    for name, indent_pt in (("TOC 1", 0), ("TOC 2", 18)):
        try:
            style = styles[name]
        except KeyError:
            style = styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
            style.base_style = base
        style.paragraph_format.left_indent = Pt(indent_pt)

def generate_toc_entries(doc, chapitres_data):
    """Génère une table des matières STATIQUE (sans champ Word) à partir des chapitres.

    Visible immédiatement à l'ouverture, sans mise à jour des champs (compatible KDP).
    """
    ensure_toc_styles(doc)
    for chap in chapitres_data:
        num = chap.get("numero", 1)
        c_titre = chap.get("titre", "").strip()
        doc.add_paragraph(f"Chapitre {num} — {c_titre}", style="TOC 1")
        for sub in chap.get("sous_chapitres", []):
            s_titre = sub.get("titre", "").strip()
            if s_titre:
                doc.add_paragraph(s_titre, style="TOC 2")

def configure_styles(doc):
    """Configure la police Georgia et les tailles demandées sur les styles de base."""
    # Style Normal
    style_normal = doc.styles['Normal']
    font_normal = style_normal.font
    font_normal.name = 'Georgia'
    font_normal.size = Pt(11)
    
    # Style Heading 1
    style_h1 = doc.styles['Heading 1']
    font_h1 = style_h1.font
    font_h1.name = 'Georgia'
    font_h1.size = Pt(16)
    font_h1.bold = True
    
    # Style Heading 2
    style_h2 = doc.styles['Heading 2']
    font_h2 = style_h2.font
    font_h2.name = 'Georgia'
    font_h2.size = Pt(13)
    font_h2.bold = True

def clean_markdown(text):
    """Nettoie les syntaxes Markdown courantes en préservant le texte brut des titres."""
    if not text:
        return ""
    # Convertit les balises HTML <br>, <br/>, <br /> (et <br><br>) en sauts de ligne
    text = re.sub(r'<br\s*/?>', '\n', text, flags=re.IGNORECASE)
    lines = text.split('\n')
    cleaned_lines = []
    for line in lines:
        stripped = line.strip()
        
        # Si c'est un titre MD (ex: ### Sous-titre interne), on garde le texte sans les #
        if stripped.startswith('#'):
            cleaned_line = stripped.lstrip('#').strip()
            cleaned_lines.append(cleaned_line)
            continue
            
        # Enlève les puces de liste MD basiques (* ou -)
        if stripped.startswith('* ') or stripped.startswith('- '):
            line = stripped[2:]
            
        # Ignore les lignes de séparation de tableaux MD
        if '|' in line and '---' in line:
            continue
            
        cleaned_lines.append(line)
    return '\n'.join(cleaned_lines).strip()

def add_paragraph_with_markdown(doc, texte, style='Normal'):
    """Ajoute du texte de chapitre en interprétant le markdown inline.

    - Nettoie le markdown (titres, puces, tableaux, <br> -> saut de ligne).
    - Crée un paragraphe par ligne non vide.
    - Applique réellement le gras **...** via des runs distincts (run.bold = True).
    """
    cleaned = clean_markdown(texte)
    if not cleaned:
        return
    for line in cleaned.split('\n'):
        if not line.strip():
            continue
        p = doc.add_paragraph(style=style)
        # Découpe la ligne en alternant segments **gras** et texte normal
        for segment in re.split(r'(\*\*.+?\*\*)', line):
            if not segment:
                continue
            if segment.startswith('**') and segment.endswith('**') and len(segment) > 4:
                run = p.add_run(segment[2:-2])
                run.bold = True
            else:
                p.add_run(segment)

def main():
    parser = argparse.ArgumentParser()
    parser.add_argument('--config', required=True, help='Path to book_config.json')
    args = parser.parse_args()
    
    try:
        # 1. Chargement de la configuration
        if not os.path.exists(args.config):
            raise FileNotFoundError(f"Fichier de configuration introuvable : {args.config}")
            
        with open(args.config, 'r', encoding='utf-8') as f:
            config = json.load(f)
            
        titre = config.get("titre_livre", "").strip()
        sous_titre = config.get("sous_titre", "").strip()
        auteur = config.get("auteur_nom", "").strip()
        bio = config.get("auteur_bio", "").strip()
        chapitres_dir = config.get("chapitres_dir", "")
        output_path = config.get("output_path", "")

        # --- 0. DIAGNOSTIC DOSSIER CHAPITRES ---
        LIVRES_ROOT = "C:/Users/luken/.n8n-files/Livres/"
        diag_pattern = os.path.join(chapitres_dir, "chapitre_*.json")

        sys.stderr.write("=== DIAGNOSTIC CHAPITRES ===\n")
        sys.stderr.write(f"chapitres_dir (config): {chapitres_dir}\n")
        sys.stderr.write(f"os.path.exists(chapitres_dir): {os.path.exists(chapitres_dir)}\n")

        dossiers_disponibles = []
        if os.path.exists(LIVRES_ROOT):
            try:
                dossiers_disponibles = os.listdir(LIVRES_ROOT)
            except Exception as e:
                sys.stderr.write(f"Erreur listing {LIVRES_ROOT}: {str(e)}\n")
        else:
            sys.stderr.write(f"ATTENTION: racine introuvable: {LIVRES_ROOT}\n")

        sys.stderr.write(f"Contenu de {LIVRES_ROOT} ({len(dossiers_disponibles)}):\n")
        for d in dossiers_disponibles:
            sys.stderr.write(f"  - {d}\n")

        sys.stderr.write(f"glob pattern: {diag_pattern}\n")
        glob_initial = glob.glob(diag_pattern)
        sys.stderr.write(f"glob.glob() -> {len(glob_initial)} fichier(s):\n")
        for f_found in glob_initial:
            sys.stderr.write(f"  - {f_found}\n")

        # --- FALLBACK: recherche d'un dossier approchant ---
        if not glob_initial:
            sys.stderr.write("Aucun chapitre via le dossier exact. Recherche d'un fallback...\n")
            fallback_dir = None
            for d in dossiers_disponibles:
                full = os.path.join(LIVRES_ROOT, d)
                if not os.path.isdir(full):
                    continue
                low = d.lower()
                if "tete" in low or "hors" in low:
                    fallback_dir = full
                    sys.stderr.write(f"Dossier fallback retenu: {full}\n")
                    break
            if fallback_dir:
                chapitres_dir = fallback_dir
            else:
                sys.stderr.write("Aucun dossier fallback approchant trouve.\n")
        sys.stderr.write("=== FIN DIAGNOSTIC ===\n")
        sys.stderr.flush()

        doc = docx.Document()
        configure_styles(doc)
        
        # --- 1. PAGE DE TITRE ---
        p_titre = doc.add_paragraph()
        p_titre.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_titre.paragraph_format.space_before = Pt(120)
        run_titre = p_titre.add_run(titre)
        run_titre.font.size = Pt(24)
        run_titre.font.bold = True
        
        if sous_titre:
            p_sub = doc.add_paragraph()
            p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_sub.paragraph_format.space_before = Pt(12)
            run_sub = p_sub.add_run(sous_titre)
            run_sub.font.size = Pt(14)
            run_sub.font.italic = True
            
        p_aut = doc.add_paragraph()
        p_aut.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_aut.paragraph_format.space_before = Pt(48)
        run_aut = p_aut.add_run(auteur)
        run_aut.font.size = Pt(13)
        
        # --- 2. SAUT DE PAGE ---
        doc.add_page_break()
        
        # --- 3. PAGE COPYRIGHT ---
        p_copy = doc.add_paragraph()
        p_copy.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_copy.paragraph_format.space_before = Pt(200)
        p_copy.add_run(f"© 2026 {auteur}. Tous droits réservés. Toute reproduction, même partielle, est interdite sans autorisation préalable de l'auteur.")
        
        # --- 4. SAUT DE PAGE ---
        doc.add_page_break()
        
        # --- 5. CHARGEMENT ET SÉLECTION DES CHAPITRES (avant la TOC) ---
        json_pattern = os.path.join(chapitres_dir, "chapitre_*.json")
        chapitre_files = glob.glob(json_pattern)

        if not chapitre_files:
            dispo = ", ".join(dossiers_disponibles) if dossiers_disponibles else "(aucun)"
            print(json.dumps({
                "status": "error",
                "message": f"Aucun chapitre trouvé. Dossiers disponibles: [{dispo}]"
            }, ensure_ascii=False))
            sys.exit(1)

        chapitres_data = []
        for file_path in chapitre_files:
            try:
                with open(file_path, 'r', encoding='utf-8') as cf:
                    c_data = json.load(cf)
                    if "numero" in c_data:
                        chapitres_data.append(c_data)
            except Exception as e:
                sys.stderr.write(f"Warning: Impossible de lire le fichier {file_path}: {str(e)}\n")

        # Tri strict par numéro
        chapitres_data.sort(key=lambda x: int(x["numero"]))

        # --- 6. TABLE DES MATIÈRES (statique, visible sans mise à jour Word) ---
        doc.add_heading("Table des matières", level=1)
        generate_toc_entries(doc, chapitres_data)

        # --- 7. SAUT DE PAGE ---
        doc.add_page_break()

        # Insertion des chapitres
        for chap in chapitres_data:
            num = chap.get("numero", 1)
            c_titre = chap.get("titre", "").strip()
            
            # Titre du Chapitre
            doc.add_heading(f"Chapitre {num} — {c_titre}", level=1)
            
            # Introduction
            if chap.get("introduction"):
                add_paragraph_with_markdown(doc, chap["introduction"], style='Normal')

            # Sous-chapitres
            for sub in chap.get("sous_chapitres", []):
                s_titre = sub.get("titre", "").strip()
                s_contenu = sub.get("contenu", "")

                if s_titre:
                    doc.add_heading(s_titre, level=2)
                if s_contenu:
                    add_paragraph_with_markdown(doc, s_contenu, style='Normal')

            # Conclusion
            if chap.get("conclusion"):
                add_paragraph_with_markdown(doc, chap["conclusion"], style='Normal')
                
            # Saut de page après chaque chapitre
            doc.add_page_break()
            
        # --- 8. À PROPOS DE L'AUTEUR ---
        doc.add_heading("À propos de l'auteur", level=1)
        if bio:
            doc.add_paragraph(bio, style='Normal')
        else:
            doc.add_paragraph(f"{auteur} est l'auteur de cet ouvrage.", style='Normal')
            
        # --- SAUVEGARDE ET EXPORT ---
        output_dir = os.path.dirname(output_path)
        if output_dir:
            os.makedirs(output_dir, exist_ok=True)
            
        doc.save(output_path)
        
        # Print JSON final succès
        print(json.dumps({"status": "ok", "output": output_path}, ensure_ascii=False))
        sys.exit(0)

    except Exception as e:
        # Print JSON final échec
        print(json.dumps({"status": "error", "message": str(e)}, ensure_ascii=False))
        sys.exit(1)

if __name__ == "__main__":
    main()